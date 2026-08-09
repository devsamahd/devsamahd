import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def add_bottom_border(paragraph, color_hex="0A2540", size_pt=1.5):
    """Adds a stylish horizontal border under a section heading."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="{int(size_pt * 8)}" w:space="4" w:color="{color_hex}"/></w:pBdr>')
    pPr.append(pBdr)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding for Word tables."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def create_resume():
    doc = docx.Document()

    # Set page margins to 0.65 inches top/bottom, 0.7 inches left/right for ATS compatibility and sleek layout
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Color Palette - Professional Corporate Navy & Teal Accent
    PRIMARY_COLOR = RGBColor(10, 37, 64)       # Deep Navy #0A2540
    SECONDARY_COLOR = RGBColor(16, 185, 129)   # Emerald Teal #10B981
    TEXT_DARK = RGBColor(34, 34, 34)           # Charcoal #222222
    TEXT_MUTED = RGBColor(100, 110, 120)       # Slate Gray #646E78

    # Set Default Normal Style Font
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(10.5)
    font_normal.color.rgb = TEXT_DARK

    # ==========================================
    # 1. HEADER SECTION (Contact Info)
    # ==========================================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    
    run_name = title_p.add_run("ABDULSAMAD ABDULSALAM")
    run_name.font.name = 'Arial'
    run_name.font.size = Pt(22)
    run_name.font.bold = True
    run_name.font.color.rgb = PRIMARY_COLOR

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(6)
    
    run_sub = sub_p.add_run("Senior Fullstack & Systems Infrastructure Engineer | AI & Web3 Architect")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(11.5)
    run_sub.font.bold = True
    run_sub.font.color.rgb = SECONDARY_COLOR

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(14)
    
    contact_text = "Abuja, Nigeria  |  devsamahd@gmail.com  |  Portfolio: https://samahd.is-a.dev\nLinkedIn: linkedin.com/in/abdulsamad-abdulsalam-746617193  |  GitHub: github.com/devsamahd"
    run_contact = contact_p.add_run(contact_text)
    run_contact.font.size = Pt(9.5)
    run_contact.font.color.rgb = TEXT_MUTED

    # Helper for adding section headings
    def add_section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text.upper())
        run.font.name = 'Arial'
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR
        add_bottom_border(p, color_hex="0A2540", size_pt=1.5)
        return p

    # ==========================================
    # 2. PROFESSIONAL SUMMARY
    # ==========================================
    add_section_heading("Professional Summary")
    
    p_sum = doc.add_paragraph()
    p_sum.paragraph_format.space_after = Pt(8)
    p_sum.paragraph_format.line_spacing = 1.15
    run_sum = p_sum.add_run(
        "Results-driven Senior Fullstack & Backend Systems Engineer with over 4 years of hands-on experience architecting high-throughput microservices, scalable distributed networks, AI RAG pipelines, and Web3 decentralized applications. Creator of Zentramesh—a high-performance network infrastructure engine powering enterprise ISP management and customer billing systems (Netra.ng and Superkonnect) serving hundreds of active users and managing millions of NGN in revenue. Expert in Go (Golang), Node.js, TypeScript, Rust, Next.js, and cloud orchestration (Docker, Redis, PostgreSQL). Proven capability to design resilient systems with high availability, low-latency performance (<50ms), and strict API security standards."
    )

    # ==========================================
    # 3. TECHNICAL SKILLS
    # ==========================================
    add_section_heading("Technical Skills")

    skills_table = doc.add_table(rows=6, cols=2)
    skills_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    skills_table.autofit = False

    # Set column widths
    for row in skills_table.rows:
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(4.8)
        set_cell_margins(row.cells[0], top=40, bottom=40, left=60, right=60)
        set_cell_margins(row.cells[1], top=40, bottom=40, left=60, right=60)

    skills_data = [
        ("Backend & Distributed Systems:", "Go (Golang), Node.js (Express, NestJS), Rust, PHP (Laravel), RESTful APIs, WebSockets, gRPC, Microservices"),
        ("AI & Intelligent Automation:", "Retrieval-Augmented Generation (RAG), OpenAI API, Vector Databases (PGVector, Pinecone), Async Queue Processing (RabbitMQ, Redis)"),
        ("Web3 & Blockchain Technology:", "TON Blockchain, Smart Contracts (Tact), Ethers.js, Web3.js, IPFS, Decentralized Escrow Protocols, Telegram Mini Apps API"),
        ("Cloud, DevOps & Databases:", "Docker, Kubernetes, AWS (EC2), Linux (Ubuntu), Nginx, Redis, PostgreSQL, MySQL, MongoDB, Kafka, BullMQ, PM2"),
        ("Frontend & Mobile UI:", "TypeScript, JavaScript, Next.js, React.js, SwiftUI, Tailwind CSS, Chakra UI, Framer Motion, HTML5/CSS3"),
        ("SDN & Network Engineering:", "TP-Link Omada SDN API Integration, ER605 Router Orchestration, EAP 650 APs, Captive Portal Access Control & Bandwidth Shaping")
    ]

    for idx, (cat, items) in enumerate(skills_data):
        row_cells = skills_table.rows[idx].cells
        
        p_cat = row_cells[0].paragraphs[0]
        p_cat.paragraph_format.space_after = Pt(1)
        r_cat = p_cat.add_run(cat)
        r_cat.font.bold = True
        r_cat.font.size = Pt(9.5)
        r_cat.font.color.rgb = PRIMARY_COLOR

        p_item = row_cells[1].paragraphs[0]
        p_item.paragraph_format.space_after = Pt(1)
        r_item = p_item.add_run(items)
        r_item.font.size = Pt(9.5)

    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(2)

    # ==========================================
    # 4. FEATURED PROJECTS & SYSTEM ARCHITECTURE
    # ==========================================
    add_section_heading("Featured Projects & System Architecture")

    projects = [
        {
            "name": "Zentramesh — Network & ISP Infrastructure Engine",
            "sub": "Powers Netra.ng & Superkonnect  |  Go, Node.js, Next.js, Omada SDN API, PostgreSQL, Redis, PM2, Docker",
            "bullets": [
                "Architected Zentramesh, a resilient network orchestration engine interfacing directly with cloud-hosted TP-Link Omada SDN controllers across ER605 enterprise routers and EAP 650 Wi-Fi access points.",
                "Engineered back-end service layers for Netra.ng (ISP Dashboard) and Superkonnect (Access & Billing Portal), serving hundreds of daily active users and managing millions of NGN in subscription revenue.",
                "Implemented automated user access control, captive portal payment verification, dynamic bandwidth allocation, real-time subscriber diagnostics, and PM2 process monitoring with zero service downtime."
            ]
        },
        {
            "name": "Tonzscrow — Decentralized Escrow Microservice",
            "sub": "Go (Golang), Tact, PostgreSQL, TON Blockchain, Docker, Ethers.js",
            "bullets": [
                "Built a decentralized Web3 escrow API in Go facilitating trustless peer-to-peer and e-commerce buyer-seller payments.",
                "Developed secure TON blockchain smart contracts using Tact to hold, validate multi-sig approvals, and release or refund funds automatically without intermediaries."
            ]
        },
        {
            "name": "AI-Powered RAG Microservice",
            "sub": "Node.js, Express, OpenAI API, Vector DB (PGVector/Pinecone), RabbitMQ, Redis",
            "bullets": [
                "Designed an asynchronous PDF ingestion microservice that chunks extensive documents, generates high-dimensional vector embeddings, and indexes them in vector storage.",
                "Integrated OpenAI models with RabbitMQ background message workers to execute context-aware RAG search queries with sub-second response times."
            ]
        },
        {
            "name": "High-Frequency Trading & Market Data Streaming Engine",
            "sub": "Rust, Next.js, WebSockets, Apache Kafka, TypeScript",
            "bullets": [
                "Built a high-concurrency market streaming service in Rust capable of pushing real-time order-book and crypto price feeds over WebSockets to 10,000+ concurrent clients with <50ms latency.",
                "Implemented Apache Kafka event pipelines for low-latency message queueing and seamless horizontal scaling during high market volatility."
            ]
        },
        {
            "name": "Tonzilla — Encrypted Telegram Mini App",
            "sub": "Next.js, TypeScript, Telegram Mini Apps API, TON Network, Tailwind CSS",
            "bullets": [
                "Engineered a Telegram mini app providing client-side AES file encryption, secure cloud upload, and instant peer sharing directly inside Telegram chat UI.",
                "Utilized TON decentralized storage protocol for immutable, encrypted file retention and censorship resistance."
            ]
        },
        {
            "name": "Morm — Golang MongoDB ODM",
            "sub": "Go (Golang), MongoDB Go Driver, Reflection API",
            "bullets": [
                "Developed an open-source Object Document Mapper (ODM) for Golang inspired by Mongoose, providing schema validation, hook middleware, and streamlined CRUD operations.",
                "Leveraged Go Reflection API to eliminate repetitive boilerplate code by 40% while maintaining native driver execution performance."
            ]
        },
        {
            "name": "Sociallio — Multi-Channel Marketing Automation Platform",
            "sub": "Next.js, Node.js, PostgreSQL, Redis, BullMQ",
            "bullets": [
                "Architected a centralized social media automation system allowing users to schedule, manage, and analyze content across multiple social platforms.",
                "Configured Redis-backed BullMQ queue processors to handle background cron scheduling, failure retries, and analytics aggregation."
            ]
        },
        {
            "name": "Dehub.io — Decentralized Video Streaming & Monetization Platform",
            "sub": "Next.js, TypeScript, IPFS, Livepeer, Ethers.js",
            "bullets": [
                "Developed a Web3 media streaming platform leveraging IPFS for decentralized content distribution and Livepeer for real-time video transcoding.",
                "Integrated Ethereum smart contracts to enable micro-transactions and direct creator tip monetization."
            ]
        }
    ]

    for proj in projects:
        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(5)
        p_title.paragraph_format.space_after = Pt(1)
        r_name = p_title.add_run(proj["name"])
        r_name.font.bold = True
        r_name.font.size = Pt(10.5)
        r_name.font.color.rgb = PRIMARY_COLOR

        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(2)
        r_sub = p_sub.add_run(proj["sub"])
        r_sub.font.italic = True
        r_sub.font.size = Pt(8.5)
        r_sub.font.color.rgb = TEXT_MUTED

        for b in proj["bullets"]:
            bp = doc.add_paragraph(style='List Bullet')
            bp.paragraph_format.space_after = Pt(1.5)
            bp.paragraph_format.line_spacing = 1.1
            r_b = bp.add_run(b)
            r_b.font.size = Pt(9.5)

    # ==========================================
    # 5. PROFESSIONAL WORK EXPERIENCE
    # ==========================================
    add_section_heading("Professional Work Experience")

    experiences = [
        {
            "role": "Freelance Software Engineer",
            "company": "Independent Client Contracts",
            "period": "Jan 2024 - Present",
            "bullets": [
                "Architected and deployed custom backend APIs, automated billing platforms, and cloud infrastructure solutions for startups and business clients.",
                "Designed and launched Zentramesh and its client ecosystems (Netra.ng & Superkonnect), expanding regional connectivity services and generating millions of NGN in revenue.",
                "Maintained 99.9% uptime across production web services while completing computer science university studies."
            ]
        },
        {
            "role": "Software Engineer",
            "company": "Haske Groups — Lagos, Nigeria",
            "period": "Nov 2022 - Jan 2024",
            "bullets": [
                "Spearheaded backend and software development for an integrated residential estate security and tenant management system.",
                "Optimized database indexing and server-side API responses, accelerating portal load times and improving user engagement by 35%.",
                "Integrated automated gate-pass access control systems with third-party payment gateway workflows."
            ]
        },
        {
            "role": "Backend Engineer (Contract)",
            "company": "Dreamax Ltd — Tel Aviv, Israel (Remote)",
            "period": "Apr 2023 - Sep 2023",
            "bullets": [
                "Architected a scalable automated marketing system and dynamic multi-step form engine for high-traffic client acquisition campaigns.",
                "Engineered high-throughput REST APIs using Next.js, Express.js, PHP, and MySQL backed by Knex and Objection.js."
            ]
        },
        {
            "role": "Backend Engineer",
            "company": "Frixx App (Now Zendmart) — Lagos, Nigeria",
            "period": "Jan 2023 - Apr 2023",
            "bullets": [
                "Engineered a resilient backend API for a developer social platform using TypeScript, Node.js, Express, and MongoDB.",
                "Implemented secure authentication, activity feeds, and real-time notification microservices."
            ]
        },
        {
            "role": "Software Developer",
            "company": "Education Online Nigeria — Ilorin, Nigeria",
            "period": "May 2020 - Nov 2021",
            "bullets": [
                "Built essential components for an e-learning platform, including an online digital library and a dynamic scholarship allocation application.",
                "Served over 5,000 active students, ensuring reliable access to educational materials."
            ]
        }
    ]

    for exp in experiences:
        p_head = doc.add_paragraph()
        p_head.paragraph_format.space_before = Pt(5)
        p_head.paragraph_format.space_after = Pt(1)
        
        r_role = p_head.add_run(exp["role"])
        r_role.font.bold = True
        r_role.font.size = Pt(10.5)
        r_role.font.color.rgb = PRIMARY_COLOR

        r_comp = p_head.add_run(f" | {exp['company']}")
        r_comp.font.bold = True
        r_comp.font.size = Pt(10)
        r_comp.font.color.rgb = TEXT_DARK

        p_period = doc.add_paragraph()
        p_period.paragraph_format.space_after = Pt(2)
        r_per = p_period.add_run(exp["period"])
        r_per.font.italic = True
        r_per.font.size = Pt(8.5)
        r_per.font.color.rgb = SECONDARY_COLOR

        for b in exp["bullets"]:
            bp = doc.add_paragraph(style='List Bullet')
            bp.paragraph_format.space_after = Pt(1.5)
            bp.paragraph_format.line_spacing = 1.1
            r_b = bp.add_run(b)
            r_b.font.size = Pt(9.5)

    # ==========================================
    # 6. EDUCATION
    # ==========================================
    add_section_heading("Education")

    p_edu = doc.add_paragraph()
    p_edu.paragraph_format.space_before = Pt(4)
    p_edu.paragraph_format.space_after = Pt(1)
    r_deg = p_edu.add_run("Bachelor of Science (B.Sc.) in Computer Science / Engineering")
    r_deg.font.bold = True
    r_deg.font.size = Pt(10.5)
    r_deg.font.color.rgb = PRIMARY_COLOR

    p_edu_sub = doc.add_paragraph()
    p_edu_sub.paragraph_format.space_after = Pt(6)
    r_edu_sub = p_edu_sub.add_run("Focus: Software Engineering, Distributed Microservices, Database Systems & Security")
    r_edu_sub.font.italic = True
    r_edu_sub.font.size = Pt(9)
    r_edu_sub.font.color.rgb = TEXT_MUTED

    # Save output docx in workspace
    output_filename = "Abdulsamad_Abdulsalam_Portfolio_Resume.docx"
    output_path = f"/Users/user/Documents/devsamahd/{output_filename}"
    doc.save(output_path)
    print(f"Successfully generated ATS-Optimized Portfolio Resume at: {output_path}")

    # Also save in artifact directory
    artifact_path = f"/Users/user/.gemini/antigravity/brain/a00b622e-39a5-4862-9e4d-6d3e245ead66/{output_filename}"
    doc.save(artifact_path)
    print(f"Successfully saved copy to artifact path: {artifact_path}")

if __name__ == "__main__":
    create_resume()
